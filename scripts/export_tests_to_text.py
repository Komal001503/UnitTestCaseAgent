#!/usr/bin/env python3
"""Export unit test cases from .py test files to structured text files.

Parses Python unittest files, extracts test classes and methods with their
docstrings, and produces a human-readable text report grouped by user story.

Usage:
    python scripts/export_tests_to_text.py [--output-dir OUTPUT_DIR] [--format {txt,docx,xlsx}]

By default, output is written to ``test_reports/`` in the repository root.
"""

import argparse
import ast
import re
import textwrap
from pathlib import Path

# Default environment label used in Excel reports; override via constant.
DEFAULT_ENVIRONMENT = "QA – Chrome"


# ---------------------------------------------------------------------------
# Argument parsing
# ---------------------------------------------------------------------------

def parse_args():
    parser = argparse.ArgumentParser(
        description="Export Python unit tests to text (or Word) files grouped by user story."
    )
    parser.add_argument(
        "--tests-dir",
        default=None,
        help="Directory containing test_*.py files (default: tests/ in repo root)",
    )
    parser.add_argument(
        "--output-dir",
        "-o",
        default=None,
        help="Directory for output files (default: test_reports/ in repo root)",
    )
    parser.add_argument(
        "--format",
        choices=["txt", "docx", "xlsx"],
        default="txt",
        help="Output format: 'txt' (plain text), 'docx' (Word document), or 'xlsx' (Excel spreadsheet). Default: txt",
    )
    return parser.parse_args()


# ---------------------------------------------------------------------------
# AST helpers – extract structured info from test files
# ---------------------------------------------------------------------------

def _extract_docstring(node) -> str:
    """Return the docstring of an AST node, or empty string."""
    ds = ast.get_docstring(node)
    return ds.strip() if ds else ""


def _extract_test_type(docstring: str) -> str:
    """Heuristically determine test type from the docstring prefix."""
    lower = docstring.lower()
    for label in ("positive", "negative", "boundary", "integration", "edge"):
        if lower.startswith(label):
            return label.capitalize()
    return "General"


_US_PATTERN = re.compile(r"US-\d+", re.IGNORECASE)

_SOURCE_STORY_PATTERN = re.compile(
    r"""^SOURCE_STORY_FILE\s*=\s*(?:"|')(.+?)(?:"|')""",
    re.MULTILINE,
)


def _extract_user_stories(text: str) -> list[str]:
    """Return a sorted unique list of user story IDs found in *text*."""
    return sorted(set(m.group().upper() for m in _US_PATTERN.finditer(text)))


def _extract_source_story_file(filepath: Path) -> str | None:
    """Extract the SOURCE_STORY_FILE module-level constant from a test file.

    Returns the value if found, or ``None``.
    """
    source = filepath.read_text(encoding="utf-8")
    m = _SOURCE_STORY_PATTERN.search(source)
    return m.group(1) if m else None


def _extract_string_literals(node: ast.AST) -> list[str]:
    """Recursively collect all string literal values from an AST node."""
    strings: list[str] = []
    for child in ast.walk(node):
        if isinstance(child, ast.Constant) and isinstance(child.value, str):
            strings.append(child.value)
    return strings


def _extract_call_args(node: ast.AST) -> list[str]:
    """Extract string arguments from service/method calls in test body.

    Only captures arguments from calls on ``self.<service>.<method>(...)``
    patterns – i.e. the actual service invocations that represent test
    input data.  Assertion calls (``self.assertXxx``) are excluded.
    """
    args: list[str] = []
    for child in ast.walk(node):
        if not isinstance(child, ast.Call):
            continue
        func = child.func
        # Match self.<service>.<method>(...) but NOT self.assert*
        if isinstance(func, ast.Attribute):
            attr_name = func.attr
            # Skip assertions and return_value assignments
            if attr_name.startswith("assert"):
                continue
            if attr_name == "return_value":
                continue
            # Only keep calls on self.<obj>.<method>
            if isinstance(func.value, ast.Attribute):
                for arg in child.args:
                    if isinstance(arg, ast.Constant) and isinstance(arg.value, str):
                        args.append(arg.value)
    return args


def _extract_mock_return_values(node: ast.AST) -> dict[str, str]:
    """Extract key-value pairs from mock return_value dicts.

    Looks for patterns like ``service.method.return_value = {"key": "val"}``
    and returns a dict of all string key-value pairs found.
    """
    pairs: dict[str, str] = {}
    for child in ast.walk(node):
        if isinstance(child, ast.Assign):
            # Look for Dict values in assignments
            if isinstance(child.value, ast.Dict):
                for key, val in zip(child.value.keys, child.value.values):
                    if (isinstance(key, ast.Constant) and isinstance(key.value, str)
                            and isinstance(val, ast.Constant) and isinstance(val.value, str)):
                        pairs[key.value] = val.value
    return pairs


def _extract_assertions(node: ast.AST) -> list[dict]:
    """Extract assertion details from a test method.

    Returns a list of dicts with keys: method, args (list of string values).
    """
    assertions: list[dict] = []
    for child in ast.walk(node):
        if isinstance(child, ast.Call) and isinstance(child.func, ast.Attribute):
            attr_name = child.func.attr
            if attr_name.startswith("assert"):
                assertion_args: list[str] = []
                for arg in child.args:
                    if isinstance(arg, ast.Constant) and isinstance(arg.value, str):
                        assertion_args.append(arg.value)
                assertions.append({
                    "method": attr_name,
                    "args": assertion_args,
                })
    return assertions


def _extract_setup_info(class_node: ast.ClassDef) -> dict:
    """Extract precondition info from the setUp method of a test class.

    Returns a dict with keys:
        mock_names   – list of attribute names assigned as MagicMock
        setUp_doc    – docstring of setUp method
    """
    mock_names: list[str] = []
    setup_doc = ""
    for item in class_node.body:
        if isinstance(item, ast.FunctionDef) and item.name == "setUp":
            setup_doc = _extract_docstring(item)
            for stmt in ast.walk(item):
                # Detect self.xxx = MagicMock()
                if isinstance(stmt, ast.Assign):
                    for target in stmt.targets:
                        if (isinstance(target, ast.Attribute)
                                and isinstance(target.value, ast.Name)
                                and target.value.id == "self"):
                            mock_names.append(target.attr)
            break
    return {"mock_names": mock_names, "setUp_doc": setup_doc}


def parse_test_file(filepath: Path) -> list[dict]:
    """Parse a single test file and return a list of test-class dicts.

    Each dict has keys:
        class_name, class_docstring, user_stories, source_file,
        source_story_file, setup_info, tests
    where ``tests`` is a list of dicts with:
        method_name, docstring, test_type, call_args, mock_returns, assertions
    """
    source = filepath.read_text(encoding="utf-8")
    tree = ast.parse(source, filename=str(filepath))

    # Also extract user stories mentioned in the module docstring.
    module_docstring = _extract_docstring(tree)
    module_stories = _extract_user_stories(module_docstring)

    # Extract the optional SOURCE_STORY_FILE constant.
    source_story_file = _extract_source_story_file(filepath)

    classes: list[dict] = []
    for node in ast.walk(tree):
        if not isinstance(node, ast.ClassDef):
            continue
        cls_doc = _extract_docstring(node)
        cls_stories = _extract_user_stories(cls_doc) or module_stories
        setup_info = _extract_setup_info(node)

        methods: list[dict] = []
        for item in node.body:
            if isinstance(item, ast.FunctionDef) and item.name.startswith("test_"):
                m_doc = _extract_docstring(item)
                methods.append({
                    "method_name": item.name,
                    "docstring": m_doc,
                    "test_type": _extract_test_type(m_doc),
                    "call_args": _extract_call_args(item),
                    "mock_returns": _extract_mock_return_values(item),
                    "assertions": _extract_assertions(item),
                })

        if methods:
            classes.append({
                "class_name": node.name,
                "class_docstring": cls_doc,
                "user_stories": cls_stories,
                "source_file": filepath.name,
                "source_story_file": source_story_file,
                "setup_info": setup_info,
                "tests": methods,
            })

    return classes


# ---------------------------------------------------------------------------
# Group parsed data by user story
# ---------------------------------------------------------------------------

def group_by_user_story(all_classes: list[dict]) -> dict[str, list[dict]]:
    """Return {story_id: [class_dict, ...]} mapping."""
    grouped: dict[str, list[dict]] = {}
    for cls in all_classes:
        for story_id in cls["user_stories"]:
            grouped.setdefault(story_id, []).append(cls)
    return dict(sorted(grouped.items(), key=lambda kv: kv[0]))


def group_by_source_story(all_classes: list[dict]) -> dict[str | None, list[dict]]:
    """Return {source_story_file: [class_dict, ...]} mapping.

    Classes whose ``source_story_file`` is ``None`` are grouped under the
    ``None`` key, which represents the default (un-attributed) bucket.
    """
    grouped: dict[str | None, list[dict]] = {}
    for cls in all_classes:
        key = cls.get("source_story_file")
        grouped.setdefault(key, []).append(cls)
    return grouped


# ---------------------------------------------------------------------------
# Rendering – plain text
# ---------------------------------------------------------------------------

def _story_sort_key(story_id: str) -> int:
    """Extract numeric part of US-NNN for sorting."""
    m = re.search(r"\d+", story_id)
    return int(m.group()) if m else 0


def render_text(grouped: dict[str, list[dict]]) -> str:
    """Render grouped test data as a human-readable plain-text report."""
    lines: list[str] = []
    lines.append("=" * 80)
    lines.append("UNIT TEST CASES REPORT")
    lines.append("Generated from Python test files")
    lines.append("=" * 80)
    lines.append("")

    sorted_stories = sorted(grouped.keys(), key=_story_sort_key)

    for story_id in sorted_stories:
        classes = grouped[story_id]
        lines.append("-" * 80)
        lines.append(f"USER STORY: {story_id}")
        lines.append("-" * 80)
        lines.append("")

        test_counter = 0
        for cls in classes:
            lines.append(f"  Test Class: {cls['class_name']}")
            if cls["class_docstring"]:
                lines.append(f"  Description: {cls['class_docstring']}")
            lines.append(f"  Source File: {cls['source_file']}")
            lines.append("")

            # Build scenario table
            lines.append(f"  {'#':<4} {'Type':<14} {'Test Method':<55} Description")
            lines.append(f"  {'─'*4} {'─'*14} {'─'*55} {'─'*40}")

            for test in cls["tests"]:
                test_counter += 1
                ttype = test["test_type"]
                name = test["method_name"]
                desc = test["docstring"] or "(no description)"
                # Remove type prefix from description since we show it separately
                desc_clean = re.sub(
                    r"^(Positive|Negative|Boundary|Integration|Edge|General)\s*:\s*",
                    "",
                    desc,
                    flags=re.IGNORECASE,
                )
                lines.append(f"  {test_counter:<4} {ttype:<14} {name:<55} {desc_clean}")

            lines.append("")

        lines.append(f"  Total test cases for {story_id}: {test_counter}")
        lines.append("")

    # Summary
    lines.append("=" * 80)
    lines.append("SUMMARY")
    lines.append("=" * 80)
    total = 0
    for story_id in sorted_stories:
        count = sum(len(c["tests"]) for c in grouped[story_id])
        total += count
        lines.append(f"  {story_id}: {count} test case(s)")
    lines.append(f"\n  TOTAL TEST CASES: {total}")
    lines.append("")

    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Rendering – Word document (docx)
# ---------------------------------------------------------------------------

def render_docx(grouped: dict[str, list[dict]], output_path: Path) -> None:
    """Render grouped test data as a Word (.docx) document."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise SystemExit(
            "python-docx is required for Word output. Install it with:\n"
            "  pip install python-docx"
        )

    doc = Document()

    # Title
    title = doc.add_heading("Unit Test Cases Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Generated from Python test files")
    doc.add_paragraph("")

    sorted_stories = sorted(grouped.keys(), key=_story_sort_key)

    for story_id in sorted_stories:
        classes = grouped[story_id]
        doc.add_heading(f"User Story: {story_id}", level=1)

        for cls in classes:
            doc.add_heading(cls["class_name"], level=2)
            if cls["class_docstring"]:
                doc.add_paragraph(cls["class_docstring"])
            doc.add_paragraph(f"Source File: {cls['source_file']}")

            # Table: #, Type, Test Method, Description
            table = doc.add_table(rows=1, cols=4)
            table.style = "Light Grid Accent 1"
            hdr = table.rows[0].cells
            hdr[0].text = "#"
            hdr[1].text = "Type"
            hdr[2].text = "Test Method"
            hdr[3].text = "Description"

            for idx, test in enumerate(cls["tests"], start=1):
                row = table.add_row().cells
                row[0].text = str(idx)
                row[1].text = test["test_type"]
                row[2].text = test["method_name"]
                desc = test["docstring"] or "(no description)"
                desc_clean = re.sub(
                    r"^(Positive|Negative|Boundary|Integration|Edge|General)\s*:\s*",
                    "",
                    desc,
                    flags=re.IGNORECASE,
                )
                row[3].text = desc_clean

            doc.add_paragraph("")

    # Summary
    doc.add_heading("Summary", level=1)
    total = 0
    for story_id in sorted_stories:
        count = sum(len(c["tests"]) for c in grouped[story_id])
        total += count
        doc.add_paragraph(f"{story_id}: {count} test case(s)")
    doc.add_paragraph(f"\nTotal Test Cases: {total}")

    doc.save(str(output_path))


# ---------------------------------------------------------------------------
# Rendering – Excel spreadsheet (xlsx)
# ---------------------------------------------------------------------------

def _build_test_case_id(story_id: str, index: int) -> str:
    """Build a test case ID from a user story ID and index.

    Example: ``_build_test_case_id("US-100", 1)`` → ``"TC_US_100_001"``
    """
    return f"TC_{story_id.replace('-', '_')}_{index:03d}"


def _extract_module_from_source(source_file: str) -> str:
    """Derive a module name from the test source filename."""
    # test_us001_us004_us007_us012_us016_us027_login.py -> Login
    name = source_file.replace("test_", "").replace(".py", "")
    # Remove user-story segments like us001, us004, etc.
    parts = re.sub(r"us\d+_?", "", name, flags=re.IGNORECASE).strip("_")
    if parts:
        return parts.replace("_", " ").title()
    return "General"


def _format_test_steps(method_name: str, docstring: str = "",
                       test_info: dict | None = None) -> str:
    """Convert a test method name and docstring into detailed numbered test steps.

    When *test_info* is provided (from the enhanced AST parser) the steps
    include concrete actions derived from the actual test code – call
    arguments, mock scenarios, and assertion checks.

    When a *docstring* is provided the steps include a human-readable
    description derived from it, making the output more meaningful than
    relying on the method name alone.
    """
    # Strip type prefix from docstring so we get a clean description.
    desc = re.sub(
        r"^(Positive|Negative|Boundary|Integration|Edge|General)\s*:\s*",
        "",
        docstring,
        flags=re.IGNORECASE,
    ).strip() if docstring else ""

    parts = method_name.split("_")
    # Remove leading "test" token
    parts = [p for p in parts if p.lower() != "test"]

    # Derive a human-readable action from method name parts
    action = parts[0] if parts else method_name
    scenario = parts[1] if len(parts) >= 2 else ""
    expected_token = "_".join(parts[2:]) if len(parts) >= 3 else ""

    # Pretty-print scenario: camelCase → spaced words
    readable_scenario = re.sub(r"([a-z])([A-Z])", r"\1 \2", scenario).lower() if scenario else ""
    readable_action = re.sub(r"([a-z])([A-Z])", r"\1 \2", action).lower() if action else ""

    steps: list[str] = []

    # Step 1: Navigate/Prepare - specific to the action context
    method_lower = method_name.lower()
    if "login" in method_lower or "sso" in method_lower:
        steps.append("1. Navigate to the application login page")
    elif "display" in method_lower or "render" in method_lower or "show" in method_lower:
        steps.append(f"1. Navigate to the {readable_action} screen and wait for it to fully load")
    elif "upload" in method_lower or "bulk" in method_lower:
        steps.append("1. Navigate to the upload/import section of the application")
    elif "profile" in method_lower:
        steps.append("1. Navigate to the user profile page")
    elif "onboarding" in method_lower:
        steps.append("1. Navigate to the onboarding module")
    elif "approval" in method_lower or "approver" in method_lower:
        steps.append("1. Navigate to the approval workflow section")
    else:
        if readable_action:
            steps.append(f"1. Navigate to the {readable_action} section of the application")
        else:
            steps.append("1. Navigate to the relevant application module")

    # Step 2: Perform the action with specific data
    if test_info and test_info.get("call_args"):
        args_display = ", ".join(f'"{a}"' for a in test_info["call_args"])
        steps.append(
            f"2. Enter the test data: {args_display} and perform the {readable_action} action"
        )
    elif test_info and test_info.get("mock_returns"):
        mock_vals = test_info["mock_returns"]
        setup_details = ", ".join(f'{k}="{v}"' for k, v in mock_vals.items())
        steps.append(
            f"2. Trigger the {readable_action} action (service configured to return: {setup_details})"
        )
    elif "display" in method_lower or "render" in method_lower:
        element = expected_token.replace("_", " ") if expected_token else readable_scenario
        # Convert camelCase to readable words
        element = re.sub(r"([a-z])([A-Z])", r"\1 \2", element).lower()
        steps.append(
            f"2. Observe the page layout and locate the {element} element"
        )
    elif readable_scenario:
        steps.append(
            f"2. Perform the {readable_action} action with {readable_scenario} scenario"
        )
    else:
        steps.append(f"2. Execute the {readable_action} operation with prepared test data")

    # Step 3: Observe / wait for response
    if "display" in method_lower or "render" in method_lower or "show" in method_lower:
        steps.append(
            "3. Verify the UI element is visible and properly rendered on the page"
        )
    else:
        steps.append(
            "3. Wait for the system to process the request and observe the response"
        )

    # Step 4: Verify expected outcome
    if desc:
        steps.append(f"4. Validate expected result: {desc}")
    elif expected_token:
        readable_expected = expected_token.replace("_", " ").lower()
        steps.append(f"4. Validate that the system {readable_expected}")
    else:
        steps.append("4. Validate the system response matches the expected behavior")

    # Step 5: Assertion-based verification (if we have assertion info)
    if test_info and test_info.get("assertions"):
        assertion_details = []
        for assertion in test_info["assertions"]:
            method = assertion.get("method", "")
            args = assertion.get("args", [])
            if method == "assertEqual" and args:
                assertion_details.append(f'value equals "{args[0]}"')
            elif method == "assertIn" and args:
                assertion_details.append(f'response contains "{args[0]}"')
            elif method == "assertTrue":
                assertion_details.append("condition is true")
            elif method == "assertFalse":
                assertion_details.append("condition is false")
            elif method == "assertIsNone":
                assertion_details.append("result is None/empty")
            elif method == "assertIsNotNone":
                assertion_details.append("result is not None/empty")
        if assertion_details:
            checks = "; ".join(assertion_details)
            steps.append(f"5. Confirm assertion checks pass: {checks}")

    return "\n".join(steps)


def _format_preconditions(class_doc: str, setup_info: dict | None = None,
                          module: str = "", test_type: str = "") -> str:
    """Generate detailed preconditions from class docstring and setUp info.

    Returns a multi-line string describing what must already be true
    before the test can be executed.
    """
    lines: list[str] = []

    # Module-specific preconditions (concrete system-state requirements)
    if module:
        mod_lower = module.lower()
        if "login" in mod_lower:
            lines.append("1. User account must be registered and active in the system")
            lines.append("2. Application login page must be accessible via browser")
            lines.append("3. Network connectivity to authentication server must be available")
            if test_type == "Negative":
                lines.append("4. Invalid/expired credentials must be prepared for testing")
            elif test_type == "Boundary":
                lines.append("4. Boundary test data (empty strings, max-length inputs) must be prepared")
            else:
                lines.append("4. Valid user credentials must be available for testing")
        elif "onboarding" in mod_lower:
            lines.append("1. User must be logged in with HR/Admin role permissions")
            lines.append("2. Onboarding module must be enabled and accessible")
            lines.append("3. Required employee data templates must be configured")
            if test_type == "Negative":
                lines.append("4. Invalid/incomplete onboarding data must be prepared")
            else:
                lines.append("4. Valid employee records must be available in the system")
        elif "profile" in mod_lower:
            lines.append("1. User must be authenticated and logged in")
            lines.append("2. User profile page must be accessible")
            lines.append("3. Profile data fields must be editable based on user role")
        elif "upload" in mod_lower or "bulk" in mod_lower:
            lines.append("1. User must be logged in with data upload permissions")
            lines.append("2. Bulk upload module must be accessible")
            lines.append("3. Valid file templates must be available for upload")
            if test_type == "Negative":
                lines.append("4. Invalid/corrupted files must be prepared for testing")
            else:
                lines.append("4. Test data files must be formatted correctly")
        elif "rehire" in mod_lower:
            lines.append("1. User must be logged in with HR permissions")
            lines.append("2. Rehire module must be accessible")
            lines.append("3. Previously separated employee records must exist in the system")
        elif "overview" in mod_lower:
            lines.append("1. User must be logged in with appropriate role")
            lines.append("2. Onboarding overview dashboard must be accessible")
            lines.append("3. Onboarding records must exist for display")
        elif "approver" in mod_lower or "ir" in mod_lower:
            lines.append("1. User must be logged in with IR Approver role")
            lines.append("2. Approval workflow module must be accessible")
            lines.append("3. Pending approval requests must exist in the system")
        elif "workforce" in mod_lower:
            lines.append("1. User must be logged in with workforce management role")
            lines.append("2. Workforce information module must be accessible")
            lines.append("3. Employee workforce records must exist in the system")
        else:
            lines.append(f"1. User must be logged in with access to the {module} module")
            lines.append(f"2. {module} module must be accessible and functional")
            lines.append("3. Required test data must be available in the system")
    else:
        lines.append("1. Application must be running and accessible")
        lines.append("2. User must be authenticated with appropriate permissions")
        lines.append("3. Required test data must be available")

    # Mock/service requirements
    if setup_info and setup_info.get("mock_names"):
        services = [name.replace("_", " ").replace("self.", "").title()
                     for name in setup_info["mock_names"]]
        lines.append(f"{len(lines) + 1}. Backend service(s) must be operational: {', '.join(services)}")

    return "\n".join(lines)


def _format_test_data(test_info: dict | None = None,
                      method_name: str = "", test_type: str = "") -> str:
    """Generate detailed test data description from extracted method info.

    Returns a multi-line string listing the concrete data values used
    in the test, making it easy for testers to replicate the scenario.
    """
    lines: list[str] = []

    # Extract call arguments from test_info
    call_args = test_info.get("call_args", []) if test_info else []
    mock_returns = test_info.get("mock_returns", {}) if test_info else {}
    assertions = test_info.get("assertions", []) if test_info else []

    # Parse method name to understand context
    parts = method_name.split("_")
    parts = [p for p in parts if p.lower() != "test"]
    action = parts[0].lower() if parts else ""

    # Call arguments – the primary input data
    if call_args:
        if "login" in action or "sso" in action.lower():
            if len(call_args) >= 2:
                lines.append(f"Username: {call_args[0]}")
                lines.append(f"Password: {call_args[1]}")
            elif len(call_args) == 1:
                if "sso" in action.lower() or "token" in method_name.lower():
                    lines.append(f"SSO Token: {call_args[0]}")
                else:
                    lines.append(f"Input: {call_args[0]}")
        else:
            for i, arg in enumerate(call_args, 1):
                lines.append(f"Input {i}: {arg}")

    # Mock return values – the expected service behavior
    if mock_returns:
        lines.append("Expected service response:")
        for key, val in mock_returns.items():
            lines.append(f"  {key}: {val}")

    # If no explicit data from call_args/mock_returns, derive data from
    # assertions and method name context
    if not call_args and not mock_returns:
        # Extract expected values from assertions
        assertion_values = []
        for assertion in assertions:
            for arg in assertion.get("args", []):
                assertion_values.append(arg)

        if assertion_values:
            lines.append("Expected UI elements/values:")
            for val in assertion_values:
                readable_val = val.replace("_", " ").title()
                lines.append(f"  - {readable_val}")
        else:
            # Derive meaningful test data from method name
            method_lower = method_name.lower()
            if "display" in method_lower or "render" in method_lower or "show" in method_lower:
                # UI verification test
                element = "_".join(parts[2:]) if len(parts) > 2 else "element"
                readable_element = element.replace("_", " ").title()
                lines.append(f"UI Element: {readable_element}")
                lines.append("Page State: Fully loaded")
            elif "valid" in method_lower or "success" in method_lower:
                lines.append("Input: Valid test data as per acceptance criteria")
            elif "invalid" in method_lower or "error" in method_lower or "fail" in method_lower:
                lines.append("Input: Invalid/malformed test data")
            elif "empty" in method_lower or "blank" in method_lower:
                lines.append("Input: Empty/blank values")
            elif "max" in method_lower or "min" in method_lower:
                lines.append("Input: Boundary value data (min/max length)")
            else:
                # Generic but still informative
                readable_action = " ".join(parts[:2]) if len(parts) >= 2 else action
                lines.append(f"Test Scenario: {readable_action.replace('_', ' ').title()}")

    # Add test type context
    if test_type == "Positive":
        lines.append("Data Type: Valid (positive test)")
    elif test_type == "Negative":
        lines.append("Data Type: Invalid (negative test - error handling)")
    elif test_type == "Boundary":
        lines.append("Data Type: Boundary/edge case values")
    elif test_type == "Integration":
        lines.append("Data Type: Integration test data")

    return "\n".join(lines)


def _format_expected_result(docstring: str = "", test_info: dict | None = None,
                            test_type: str = "") -> str:
    """Generate a detailed expected result description.

    Combines the test docstring, assertion details, and mock return values
    to produce a comprehensive expected-outcome description.
    """
    # Clean description from docstring
    desc = re.sub(
        r"^(Positive|Negative|Boundary|Integration|Edge|General)\s*:\s*",
        "",
        docstring,
        flags=re.IGNORECASE,
    ).strip() if docstring else ""

    lines: list[str] = []

    if desc:
        lines.append(desc)

    if test_info:
        mock_returns = test_info.get("mock_returns", {})
        assertions = test_info.get("assertions", [])

        # Add specifics from mock return values
        if "status" in mock_returns:
            status_val = mock_returns["status"]
            if status_val == "success":
                lines.append("The system should process the request successfully")
            elif status_val == "error":
                lines.append("The system should return an appropriate error response")

        if "redirect" in mock_returns:
            lines.append(f"User should be redirected to: {mock_returns['redirect']}")

        if "message" in mock_returns:
            lines.append(f"System should display message: \"{mock_returns['message']}\"")

        # Add assertion-based expectations
        for assertion in assertions:
            method = assertion.get("method", "")
            args = assertion.get("args", [])
            if method == "assertEqual" and args:
                for arg in args:
                    lines.append(f"Verify value equals: \"{arg}\"")
            elif method == "assertIn" and args:
                for arg in args:
                    lines.append(f"Verify response contains: \"{arg}\"")

    # Add general expectation based on test type
    if test_type == "Positive":
        lines.append("No error messages should be displayed")
    elif test_type == "Negative":
        lines.append("An appropriate error message should be shown to the user")
    elif test_type == "Boundary":
        lines.append("System should handle the edge case gracefully without crashing")

    # Deduplicate while preserving order
    seen: set[str] = set()
    unique_lines: list[str] = []
    for line in lines:
        if line not in seen:
            seen.add(line)
            unique_lines.append(line)

    return "\n".join(unique_lines)


def render_xlsx(grouped: dict[str, list[dict]], output_path: Path) -> None:
    """Render grouped test data as an Excel (.xlsx) spreadsheet.

    Columns follow the standard test case format:
    Test Case ID, Test Case Title, Module, User Story ID, Preconditions,
    Test Data, Test Steps, Expected Result, Actual Result, Status,
    Priority, Severity, Environment, Remarks
    """
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    except ImportError:
        raise SystemExit(
            "openpyxl is required for Excel output. Install it with:\n"
            "  pip install openpyxl"
        )

    wb = Workbook()
    ws = wb.active
    ws.title = "Test Cases"

    # Define headers
    headers = [
        "Test Case ID",
        "Test Case Title",
        "Test Type",
        "Module",
        "User Story ID",
        "Preconditions",
        "Test Data",
        "Test Steps",
        "Expected Result",
        "Actual Result",
        "Status",
        "Priority",
        "Severity",
        "Environment",
        "Remarks",
    ]

    # Style for header row
    header_font = Font(bold=True, color="FFFFFF", size=11)
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    thin_border = Border(
        left=Side(style="thin"),
        right=Side(style="thin"),
        top=Side(style="thin"),
        bottom=Side(style="thin"),
    )

    # Write headers
    for col_idx, header in enumerate(headers, start=1):
        cell = ws.cell(row=1, column=col_idx, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_alignment
        cell.border = thin_border

    # Write data rows
    sorted_stories = sorted(grouped.keys(), key=_story_sort_key)
    row_idx = 2
    cell_alignment = Alignment(vertical="top", wrap_text=True)

    for story_id in sorted_stories:
        classes = grouped[story_id]
        tc_counter = 0

        for cls in classes:
            module = _extract_module_from_source(cls["source_file"])
            class_doc = cls["class_docstring"]
            setup_info = cls.get("setup_info")

            for test in cls["tests"]:
                tc_counter += 1
                tc_id = _build_test_case_id(story_id, tc_counter)

                # Build title from docstring
                doc = test["docstring"] or test["method_name"]
                title = re.sub(
                    r"^(Positive|Negative|Boundary|Integration|Edge|General)\s*:\s*",
                    "",
                    doc,
                    flags=re.IGNORECASE,
                ).strip()

                test_type = test["test_type"]

                # Build test_info dict from enhanced AST data
                test_info = {
                    "call_args": test.get("call_args", []),
                    "mock_returns": test.get("mock_returns", {}),
                    "assertions": test.get("assertions", []),
                }

                # Generate detailed content for each column
                preconditions = _format_preconditions(
                    class_doc, setup_info, module, test_type
                )
                test_data = _format_test_data(
                    test_info, test["method_name"], test_type
                )
                test_steps = _format_test_steps(
                    test["method_name"], test["docstring"], test_info
                )
                expected_result = _format_expected_result(
                    test["docstring"], test_info, test_type
                )

                row_data = [
                    tc_id,                              # Test Case ID
                    title,                              # Test Case Title
                    test_type,                          # Test Type
                    module,                             # Module
                    story_id,                           # User Story ID
                    preconditions,                      # Preconditions
                    test_data,                          # Test Data
                    test_steps,                         # Test Steps
                    expected_result,                    # Expected Result
                    "To be filled during ITQA",         # Actual Result
                    "To be filled during ITQA",         # Status
                    "To be filled during ITQA",         # Priority
                    "To be filled during ITQA",         # Severity
                    DEFAULT_ENVIRONMENT,                # Environment
                    "To be filled during ITQA",         # Remarks
                ]

                for col_idx, value in enumerate(row_data, start=1):
                    cell = ws.cell(row=row_idx, column=col_idx, value=value)
                    cell.alignment = cell_alignment
                    cell.border = thin_border

                row_idx += 1

    # Auto-adjust column widths
    for col_idx, header in enumerate(headers, start=1):
        max_len = len(header)
        for row in range(2, row_idx):
            cell_value = ws.cell(row=row, column=col_idx).value
            if cell_value:
                max_len = max(max_len, min(len(str(cell_value)), 50))
        ws.column_dimensions[ws.cell(row=1, column=col_idx).column_letter].width = max_len + 4

    # Freeze header row
    ws.freeze_panes = "A2"

    wb.save(str(output_path))


def _report_basename(source_story_file: str | None) -> str:
    """Derive a report base name from the source user-story filename.

    If *source_story_file* is ``None``, returns the legacy default name
    ``"unit_test_cases_report"``.

    Otherwise the extension is stripped from the source filename and a
    ``_test_report`` suffix is appended.
    Example: ``"MaintainShiftSchedulingPlanUserStory.xlsx"``
             → ``"MaintainShiftSchedulingPlanUserStory_test_report"``
    """
    if source_story_file is None:
        return "unit_test_cases_report"
    return Path(source_story_file).stem + "_test_report"


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    args = parse_args()

    repo_root = Path(__file__).resolve().parent.parent
    tests_dir = Path(args.tests_dir) if args.tests_dir else repo_root / "tests"
    output_dir = Path(args.output_dir) if args.output_dir else repo_root / "test_reports"

    if not tests_dir.is_dir():
        raise SystemExit(f"Tests directory not found: {tests_dir}")

    output_dir.mkdir(parents=True, exist_ok=True)

    # Collect all test files (exclude __init__.py and non-test files)
    test_files = sorted(tests_dir.glob("test_*.py"))
    if not test_files:
        raise SystemExit(f"No test_*.py files found in {tests_dir}")

    # Parse all files
    all_classes: list[dict] = []
    for tf in test_files:
        all_classes.extend(parse_test_file(tf))

    # Group test classes by their source user-story file so that each
    # user-story file gets its own report.
    story_groups = group_by_source_story(all_classes)

    for source_story_file, classes_for_story in story_groups.items():
        grouped = group_by_user_story(classes_for_story)
        basename = _report_basename(source_story_file)

        if args.format == "txt":
            report = render_text(grouped)
            out_file = output_dir / f"{basename}.txt"
            out_file.write_text(report, encoding="utf-8")
            print(f"Text report written to {out_file}")
        elif args.format == "docx":
            out_file = output_dir / f"{basename}.docx"
            render_docx(grouped, out_file)
            print(f"Word document written to {out_file}")
        elif args.format == "xlsx":
            out_file = output_dir / f"{basename}.xlsx"
            render_xlsx(grouped, out_file)
            print(f"Excel spreadsheet written to {out_file}")


if __name__ == "__main__":
    main()
